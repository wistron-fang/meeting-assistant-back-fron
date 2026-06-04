# Copyright © 2026 广州金元信息科技有限公司 版权所有
# 未经授权，禁止转售或仿制。

"""会议纪要 API"""
import os
import io
import re
import shutil
import zipfile
from datetime import datetime
from typing import Optional, List


_ILLEGAL_FN_CHARS = re.compile(r'[\\/:*?"<>|\x00-\x1f]')


def _safe_filename(title: str, fallback: str = "会议纪要", max_len: int = 80) -> str:
    """文件名安全处理：去非法字符、去首尾空格/点、截长、空兜底。"""
    if not title:
        return fallback
    name = _ILLEGAL_FN_CHARS.sub("_", title)
    name = name.strip(" .\t\r\n")
    if len(name) > max_len:
        name = name[:max_len].rstrip(" .")
    return name or fallback

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form, status
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from sqlalchemy import func
from sqlalchemy.orm import Session

from core.database import get_db
from models.user import User
from models.meeting import MeetingTask
from router.auth_router import get_current_user_required
from service.meeting.tasks import process_meeting
from service.quota import (
    effective_max_concurrent,
    effective_monthly_task_quota,
    effective_monthly_token_quota,
    current_month_start,
    current_month_tokens_used,
    star_to_priority,
)
from core.celery_app import celery_app

router = APIRouter(prefix="/meeting", tags=["会议纪要"])

MAX_FILES_PER_REQUEST = 2
ALLOWED_AUDIO_EXTS = {".m4a", ".mp4", ".wav", ".aac", ".mp3"}
ALLOWED_DOC_EXTS = {".txt", ".docx"}
MAX_AUDIO_SIZE_MB = 500
MAX_DOC_SIZE_MB = 50


class MeetingTaskOut(BaseModel):
    id: int
    title: Optional[str]
    source_type: Optional[str]
    status: str
    progress: int
    stage: Optional[str]
    has_md: bool
    has_pdf: bool
    has_docx: bool
    has_transcript: bool
    error: Optional[str]
    created_at: datetime
    updated_at: datetime
    started_at: Optional[datetime] = None
    finished_at: Optional[datetime] = None

    class Config:
        from_attributes = True


def _to_out(t: MeetingTask) -> dict:
    return {
        "id": t.id,
        "title": t.title,
        "source_type": t.source_type,
        "status": t.status,
        "progress": t.progress or 0,
        "stage": t.stage,
        "has_md": bool(t.md_path and os.path.exists(t.md_path)),
        "has_pdf": bool(t.pdf_path and os.path.exists(t.pdf_path)),
        "has_docx": bool(t.docx_path and os.path.exists(t.docx_path)),
        "has_transcript": bool(t.transcript_path and os.path.exists(t.transcript_path)),
        "error": t.error,
        "created_at": t.created_at,
        "updated_at": t.updated_at,
        "started_at": t.started_at,
        "finished_at": t.finished_at,
    }


class MeetingStats(BaseModel):
    pending: int = 0
    running: int = 0
    done: int = 0
    failed: int = 0


class MeetingTaskListOut(BaseModel):
    items: List[MeetingTaskOut]
    total: int
    stats: MeetingStats


class BatchZipIn(BaseModel):
    ids: List[int]
    format: str = "all"  # md | pdf | docx | all（"both" 作为旧值兼容，等同于 "all"）

@router.post("/tasks", response_model=List[MeetingTaskOut], status_code=status.HTTP_201_CREATED)
async def create_tasks(
    title: Optional[str] = Form(None),
    text: Optional[str] = Form(None),
    files: Optional[List[UploadFile]] = File(default=None),
    audio: Optional[UploadFile] = File(default=None),
    current_user: User = Depends(get_current_user_required),
    db: Session = Depends(get_db),
):
    """
    创建会议纪要任务。
    - 粘贴模式: 传 text,可选 title
    - 上传模式: 传 files (≤2 个 .txt)
    - 录音模式: 传 audio (单文件,.m4a/.mp4/.wav/.aac/.mp3,≤500MB)
    - 并发 / 月任务数 / 月 token 三重配额按星级限制（见 service/quota.py）
    """
    has_text = text and text.strip()
    has_files = files and any(f.filename for f in files)
    has_audio = audio is not None and audio.filename
    if not has_text and not has_files and not has_audio:
        raise HTTPException(400, "请提供会议原文 (粘贴文本 / 上传 .txt / 上传录音)")

    # 三种模式互斥,避免歧义
    modes_count = sum([bool(has_text), bool(has_files), bool(has_audio)])
    if modes_count > 1:
        raise HTTPException(400, "粘贴 / 上传 txt / 上传录音 三种模式只能选其一")

    if has_files and len(files) > MAX_FILES_PER_REQUEST:
        raise HTTPException(400, f"最多同时上传 {MAX_FILES_PER_REQUEST} 个文件")

    if has_audio:
        ext = os.path.splitext(audio.filename or "")[1].lower()
        if ext not in ALLOWED_AUDIO_EXTS:
            raise HTTPException(400, f"音频格式仅支持 {', '.join(sorted(ALLOWED_AUDIO_EXTS))}")

    new_cnt = len(files) if has_files else 1

    # ① 并发上限（按星级，可被用户列覆盖）
    max_conc = effective_max_concurrent(current_user)
    running_cnt = db.query(MeetingTask).filter(
        MeetingTask.user_id == current_user.id,
        MeetingTask.status.in_(["pending", "running"])
    ).count()
    if running_cnt + new_cnt > max_conc:
        raise HTTPException(
            429,
            f"并发上限 {max_conc},当前已有 {running_cnt} 个任务在跑,请稍候"
        )

    # ② 本月任务数配额（实时 count 本月已建任务）
    task_quota = effective_monthly_task_quota(current_user)
    month_task_cnt = db.query(MeetingTask).filter(
        MeetingTask.user_id == current_user.id,
        MeetingTask.created_at >= current_month_start(),
    ).count()
    if month_task_cnt + new_cnt > task_quota:
        raise HTTPException(
            429,
            f"本月任务配额已用完（{task_quota} 个/月,本月已建 {month_task_cnt} 个）"
        )

    # ③ 本月 token 配额（软拦截：token 是事后量，只拦「已超额」，不预扣；≤0 表示不限）
    token_quota = effective_monthly_token_quota(current_user)
    if token_quota and token_quota > 0:
        used = current_month_tokens_used(current_user)
        if used >= token_quota:
            raise HTTPException(
                429,
                f"本月 token 配额已用完（{token_quota}/月,已用 {used}）"
            )

    # 落库 + 投递
    created: List[MeetingTask] = []
    if has_audio:
        # 录音模式: 先落库拿到 task_id,再存音频到任务目录
        from service.meeting.tasks import MEETING_DATA_ROOT
        ext = os.path.splitext(audio.filename)[1].lower()
        base_title = title or os.path.splitext(os.path.basename(audio.filename))[0] or "录音会议"

        t = MeetingTask(
            user_id=current_user.id,
            title=base_title,
            source_type="audio",
            source_text=None,
            audio_filename=audio.filename,
            status="pending",
        )
        db.add(t); db.flush()

        task_dir = os.path.join(MEETING_DATA_ROOT, str(t.user_id), str(t.id))
        os.makedirs(task_dir, exist_ok=True)
        audio_path = os.path.join(task_dir, f"source{ext}")

        # 流式写入并校验大小
        size = 0
        with open(audio_path, "wb") as out:
            while chunk := await audio.read(1024 * 1024):
                size += len(chunk)
                if size > MAX_AUDIO_SIZE_MB * 1024 * 1024:
                    out.close()
                    os.remove(audio_path)
                    db.rollback()
                    raise HTTPException(413, f"音频超过 {MAX_AUDIO_SIZE_MB}MB 上限")
                out.write(chunk)

        t.audio_path = audio_path
        t.output_dir = task_dir
        created.append(t)
    elif has_files:
        for f in files:
            ext = os.path.splitext(f.filename or "")[1].lower()
            if ext not in ALLOWED_DOC_EXTS:
                raise HTTPException(
                    400,
                    f"文件 {f.filename or ''} 格式不支持：仅支持 {', '.join(sorted(ALLOWED_DOC_EXTS))}"
                )

            raw = await f.read()
            if len(raw) > MAX_DOC_SIZE_MB * 1024 * 1024:
                raise HTTPException(
                    413,
                    f"文件 {f.filename or ''} 超过 {MAX_DOC_SIZE_MB}MB 上限"
                )

            try:
                if ext == ".docx":
                    content = _extract_docx_text(raw)
                else:  # .txt
                    content = _decode_bytes(raw)
            except HTTPException:
                raise
            except Exception as e:
                # 解析失败立即 400，不入库，不调度
                raise HTTPException(400, f"无法解析文件 {f.filename or ''}：{e}")

            if not content or not content.strip():
                raise HTTPException(400, f"文件 {f.filename or ''} 内容为空，无法生成纪要")

            t = MeetingTask(
                user_id=current_user.id,
                # 用户没填标题：先占位 "待生成…"，任务跑完后由引擎产出的最终标题覆盖
                title=title or "标题待生成…",
                source_type="upload",
                source_text=content,
                status="pending",
            )
            db.add(t); db.flush()
            created.append(t)
    else:
        t = MeetingTask(
            user_id=current_user.id,
            # 用户没填标题：先占位 "待生成…"，任务跑完后由引擎产出的最终标题覆盖
            title=title or "标题待生成…",
            source_type="paste",
            source_text=text,
            status="pending",
        )
        db.add(t); db.flush()
        created.append(t)

    db.commit()
    priority = star_to_priority(current_user.star_level)
    for t in created:
        process_meeting.apply_async(args=[t.id], priority=priority)

    return [_to_out(t) for t in created]


def _decode_bytes(raw: bytes) -> str:
    """txt 文件解码：utf-8 优先，失败降级到 gbk，再失败用 errors='replace'"""
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_docx_text(raw: bytes) -> str:
    """
    从 .docx 字节流中提取纯文本：遍历段落 + 表格单元格，忽略图片。
    抛出的任何异常会被上层捕获后转 400。
    """
    from docx import Document  # python-docx 已在 requirements
    doc = Document(io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for tb in doc.tables:
        for row in tb.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" ".join(cells))
    return "\n\n".join(parts)


@router.get("/tasks", response_model=MeetingTaskListOut)
def list_tasks(
    status_filter: Optional[str] = Query(None, alias="status"),
    title: Optional[str] = Query(None, description="按标题模糊搜索（不区分大小写）"),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    current_user: User = Depends(get_current_user_required),
    db: Session = Depends(get_db),
):
    q = db.query(MeetingTask).filter(MeetingTask.user_id == current_user.id)
    if status_filter:
        # 支持单值或逗号分隔多值：?status=running 或 ?status=pending,running
        status_list = [s.strip() for s in status_filter.split(",") if s.strip()]
        if len(status_list) == 1:
            q = q.filter(MeetingTask.status == status_list[0])
        elif len(status_list) > 1:
            q = q.filter(MeetingTask.status.in_(status_list))
    if title and title.strip():
        # 转义 LIKE 元字符 % 和 _，避免用户输入被当成通配
        kw = title.strip().replace("\\", "\\\\").replace("%", "\\%").replace("_", "\\_")
        q = q.filter(MeetingTask.title.ilike(f"%{kw}%", escape="\\"))
    total = q.count()
    items = (q.order_by(MeetingTask.created_at.desc())
              .offset((page - 1) * page_size)
              .limit(page_size).all())

    # 全局状态统计：在与列表完全相同的筛选条件下分组计数（搜索 / status_filter 都会联动）
    stats = {"pending": 0, "running": 0, "done": 0, "failed": 0}
    rows = (q.with_entities(MeetingTask.status, func.count(MeetingTask.id))
              .group_by(MeetingTask.status).all())
    for st, cnt in rows:
        if st in stats:
            stats[st] = cnt

    return {"items": [_to_out(t) for t in items], "total": total, "stats": stats}


@router.get("/tasks/{task_id}", response_model=MeetingTaskOut)
def get_task(
    task_id: int,
    current_user: User = Depends(get_current_user_required),
    db: Session = Depends(get_db),
):
    t = db.query(MeetingTask).filter_by(id=task_id, user_id=current_user.id).first()
    if not t:
        raise HTTPException(404, "任务不存在")
    return _to_out(t)

@router.get("/tasks/{task_id}/download")
def download_task(
    task_id: int,
    format: str = Query("md", pattern="^(md|pdf|docx|transcript)$"),
    current_user: User = Depends(get_current_user_required),
    db: Session = Depends(get_db),
):
    t = db.query(MeetingTask).filter_by(id=task_id, user_id=current_user.id).first()
    if not t:
        raise HTTPException(404, "任务不存在")
    # transcript 在 ASR 完成后就可用,不强求 done
    if format != "transcript" and t.status != "done":
        raise HTTPException(400, "任务尚未完成")

    if format == "md":
        path, media = t.md_path, "text/markdown; charset=utf-8"
    elif format == "pdf":
        path, media = t.pdf_path, "application/pdf"
    elif format == "docx":
        path, media = t.docx_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:  # transcript
        path, media = t.transcript_path, "text/plain; charset=utf-8"

    if not path or not os.path.exists(path):
        raise HTTPException(404, f"{format} 文件不存在")

    # 下载文件名用任务标题 + 对应扩展名（老任务的磁盘文件名是时间戳的，下载时统一改成标题命名）
    ext_map = {"md": ".md", "pdf": ".pdf", "docx": ".docx", "transcript": ".txt"}
    download_name = _safe_filename(t.title or "") + ext_map[format]
    return FileResponse(path, media_type=media, filename=download_name)


@router.post("/tasks/batch-zip")
def batch_zip(
    body: BatchZipIn,
    current_user: User = Depends(get_current_user_required),
    db: Session = Depends(get_db),
):
    if not body.ids:
        raise HTTPException(400, "请选择至少一个任务")
    if body.format not in ("md", "pdf", "docx", "all", "both"):
        raise HTTPException(400, "format 必须是 md/pdf/docx/all")

    # "both" 作为旧值兼容，与 "all" 等同（三种格式全打包）
    fmt = "all" if body.format == "both" else body.format

    tasks = db.query(MeetingTask).filter(
        MeetingTask.id.in_(body.ids),
        MeetingTask.user_id == current_user.id,
        MeetingTask.status == "done",
    ).all()
    if not tasks:
        raise HTTPException(404, "没有可下载的已完成任务")

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for t in tasks:
            base = (t.title or f"task_{t.id}").replace("/", "_").replace("\\", "_")
            if fmt in ("md", "all") and t.md_path and os.path.exists(t.md_path):
                zf.write(t.md_path, arcname=f"{base}/{os.path.basename(t.md_path)}")
            if fmt in ("pdf", "all") and t.pdf_path and os.path.exists(t.pdf_path):
                zf.write(t.pdf_path, arcname=f"{base}/{os.path.basename(t.pdf_path)}")
            if fmt in ("docx", "all") and t.docx_path and os.path.exists(t.docx_path):
                zf.write(t.docx_path, arcname=f"{base}/{os.path.basename(t.docx_path)}")
    buf.seek(0)

    fname = f"meeting_minutes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@router.delete("/tasks/{task_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_task(
    task_id: int,
    current_user: User = Depends(get_current_user_required),
    db: Session = Depends(get_db),
):
    t = db.query(MeetingTask).filter_by(id=task_id, user_id=current_user.id).first()
    if not t:
        raise HTTPException(404, "任务不存在")

    # 撤销 celery 任务（如果还在跑）
    if t.status in ("pending", "running") and t.celery_id:
        try:
            celery_app.control.revoke(t.celery_id, terminate=True)
        except Exception:
            pass

    # 删产物
    if t.output_dir and os.path.isdir(t.output_dir):
        shutil.rmtree(t.output_dir, ignore_errors=True)

    db.delete(t)
    db.commit()
    return