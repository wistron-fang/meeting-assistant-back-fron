#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
批量会议纪要生成工具

从指定文件夹读取所有 .docx 文件（会议录音转写稿），
批量生成会议纪要（.md）和 RAG 文档（.json）。

用法：
    python batch_meeting.py ./会议录音文件夹 ./输出目录
    python batch_meeting.py ./会议录音文件夹                  # 输出到 ./output
    python batch_meeting.py                                   # 默认读 ./input，输出到 ./output

依赖：
    pip install python-docx --break-system-packages

输出结构：
    输出目录/
    ├── 第一次会议纪要/
    │   ├── 会议名称.md            ← 给人看的纪要
    │   ├── 会议名称_rag.json      ← 给向量库的RAG文档块
    │   └── 会议名称_meta.json     ← 元信息
    ├── 第二次会议纪要/
    │   ├── ...
    ├── batch_summary.json          ← 批量处理汇总报告
    └── all_rag_documents.json      ← 所有会议的RAG文档合并（可直接入库）
"""

import os
import sys
import json
import re
import time
import logging
import traceback
from datetime import datetime
from typing import List, Dict, Optional

# ============================================================
# 降级判定阈值
# ============================================================
DEGRADED_QUALITY_THRESHOLD = 85   # 质量分低于此值视为降级
DEGRADED_REQUIRE_TOPICS = True    # 主题数为 0 视为降级
DEGRADED_MIN_OUTPUT_CHARS = 200   # 输出字符少于此值视为降级

# ============================================================
# 日志初始化（UTF-8，解决 Windows 下中文乱码）
# ============================================================

def setup_logging(output_dir: str) -> str:
    """
    同时输出到控制台和 UTF-8 日志文件。
    劫持 print 转发到 logging，让所有现有 print(...) 自动写入 UTF-8 日志，
    不需要修改任何 meeting_minutes_no_speaker.py 里的 print 调用。
    """
    os.makedirs(output_dir, exist_ok=True)
    log_path = os.path.join(
        output_dir,
        f"batch_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
    )

    root = logging.getLogger()
    root.setLevel(logging.INFO)
    for h in list(root.handlers):
        root.removeHandler(h)

    # 文件 handler：强制 UTF-8
    fh = logging.FileHandler(log_path, mode='w', encoding='utf-8')
    fh.setFormatter(logging.Formatter(
        "%(asctime)s [%(levelname)s] %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S"
    ))
    root.addHandler(fh)

    # 控制台 handler：简洁格式
    ch = logging.StreamHandler(sys.stdout)
    ch.setFormatter(logging.Formatter("%(message)s"))
    root.addHandler(ch)

    # 劫持 print → 转发到 logging
    import builtins
    def _log_print(*args, **kwargs):
        sep = kwargs.get("sep", " ")
        msg = sep.join(str(a) for a in args)
        logging.info(msg)
    builtins.print = _log_print

    logging.info(f"日志文件: {log_path}")
    return log_path


# ============================================================
# docx 文本提取
# ============================================================

def load_docx_text(docx_path: str) -> str:
    """
    从 .docx 文件提取全部文本内容。
    支持：普通段落、表格内文字、多级标题。
    """
    try:
        from docx import Document
    except ImportError:
        print("=" * 60)
        print("错误：需要安装 python-docx")
        print("请运行：pip install python-docx --break-system-packages")
        print("=" * 60)
        sys.exit(1)

    doc = Document(docx_path)
    parts = []

    # 提取所有段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 提取表格中的文字（有些转写稿放在表格里）
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                parts.append(" ".join(row_texts))

    full_text = "\n\n".join(parts)

    if not full_text.strip():
        print(f"  ⚠ 警告：{os.path.basename(docx_path)} 提取到的文本为空")

    return full_text


def load_text_file(file_path: str) -> str:
    """读取纯文本文件（.txt / .md）"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def load_file_text(file_path: str) -> str:
    """根据文件扩展名自动选择读取方式"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.docx':
        return load_docx_text(file_path)
    elif ext in ('.txt', '.md', '.text'):
        return load_text_file(file_path)
    else:
        print(f"  ⚠ 不支持的文件格式: {ext}，尝试按文本读取")
        return load_text_file(file_path)


# ============================================================
# 批量处理核心
# ============================================================

def extract_date_from_filename(filename: str) -> str:
    """
    从文件名中提取日期。
    
    支持的格式：
        2023年_02月20日_xxx.docx   → 2023年2月20日
        2023年02月20日_xxx.docx    → 2023年2月20日
        2023_02_20_xxx.docx        → 2023年2月20日
        2023-02-20_xxx.docx        → 2023年2月20日
        20230220_xxx.docx          → 2023年2月20日
        02月20日_xxx.docx          → 2月20日（无年份）
    
    提取不到返回空字符串。
    """
    name = os.path.splitext(os.path.basename(filename))[0]
    
    # 模式1: 2023年_02月20日 / 2023年02月20日
    m = re.search(r'(\d{4})\D*(\d{1,2})\D*月\D*(\d{1,2})\D*日', name)
    if m:
        return f"{m.group(1)}年{int(m.group(2))}月{int(m.group(3))}日"
    
    # 模式2: 2023_02_20 / 2023-02-20
    m = re.search(r'(\d{4})[_\-](\d{1,2})[_\-](\d{1,2})', name)
    if m:
        return f"{m.group(1)}年{int(m.group(2))}月{int(m.group(3))}日"
    
    # 模式3: 20230220（连续8位数字）
    m = re.search(r'(\d{4})(\d{2})(\d{2})', name)
    if m:
        y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if 2000 <= y <= 2099 and 1 <= mo <= 12 and 1 <= d <= 31:
            return f"{y}年{mo}月{d}日"
    
    # 模式4: 02月20日（无年份）
    m = re.search(r'(\d{1,2})月(\d{1,2})日', name)
    if m:
        return f"{int(m.group(1))}月{int(m.group(2))}日"
    
    return ""


def find_meeting_files(input_dir: str, extensions: tuple = ('.docx', '.txt', '.md')) -> List[str]:
    """扫描目录，找到所有会议文件"""
    files = []
    for fname in sorted(os.listdir(input_dir)):
        if fname.startswith('~$') or fname.startswith('.'):
            continue  # 跳过临时文件和隐藏文件
        if fname.lower().endswith(extensions):
            files.append(os.path.join(input_dir, fname))
    return files


def safe_filename(name: str, max_len: int = 60) -> str:
    """将字符串转为安全的文件/目录名"""
    # 移除不安全字符
    safe = name.replace('/', '_').replace('\\', '_').replace(':', '：')
    safe = safe.replace('*', '').replace('?', '').replace('"', '')
    safe = safe.replace('<', '').replace('>', '').replace('|', '')
    safe = safe.strip('. ')
    if len(safe) > max_len:
        safe = safe[:max_len]
    return safe or "未命名会议"


def batch_process(input_dir: str, output_dir: str = "./output",
                  generate_rag: bool = True,
                  merge_rag: bool = True) -> Dict:
    """
    批量处理文件夹中的所有会议文件。

    参数：
        input_dir:    输入文件夹路径（包含 .docx / .txt 文件）
        output_dir:   输出根目录
        generate_rag: 是否生成 RAG 文档（默认 True）
        merge_rag:    是否合并所有 RAG 文档到一个 JSON（默认 True）

    返回：
        批量处理汇总信息
    """
    # 导入会议纪要系统
    from meeting_minutes_no_speaker import (
        AdvancedMeetingMinutesGenerator,
        extract_metadata_from_text,
        save_result,
    )

    # 扫描文件
    files = find_meeting_files(input_dir)
    if not files:
        print(f"未找到会议文件，请检查目录: {input_dir}")
        print(f"支持的格式: .docx, .txt, .md")
        return {"success": 0, "failed": 0, "files": []}

    print("=" * 70)
    print(f"批量会议纪要生成")
    print(f"=" * 70)
    print(f"输入目录:  {os.path.abspath(input_dir)}")
    print(f"输出目录:  {os.path.abspath(output_dir)}")
    print(f"文件数量:  {len(files)}")
    print(f"生成RAG:   {'是' if generate_rag else '否'}")
    print(f"合并RAG:   {'是' if merge_rag else '否'}")
    print(f"开始时间:  {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 70)

    for i, f in enumerate(files):
        print(f"  [{i+1}] {os.path.basename(f)}")
    print()

    # 初始化生成器（只初始化一次，复用 workflow）
    generator = AdvancedMeetingMinutesGenerator()

    os.makedirs(output_dir, exist_ok=True)

    # ===== 扫描已完成的文件（用于跳过） =====
    done_source_files = set()
    if os.path.exists(output_dir):
        for sub in os.listdir(output_dir):
            sub_path = os.path.join(output_dir, sub)
            if not os.path.isdir(sub_path):
                continue
            for fname in os.listdir(sub_path):
                if fname.endswith("_meta.json"):
                    try:
                        with open(os.path.join(sub_path, fname), 'r', encoding='utf-8') as mf:
                            meta = json.load(mf)
                            src = meta.get("source_file", "")
                            if src:
                                done_source_files.add(src)
                    except Exception:
                        pass

    if done_source_files:
        print(f"\n📋 检测到 {len(done_source_files)} 个已处理文件，将自动跳过")

    # 处理结果收集
    results_summary = []
    all_rag_docs = []  # 合并用
    success_count = 0
    fail_count = 0
    skip_count = 0
    degraded_count = 0   # 新增：降级统计
    total_start = time.time()

    for idx, file_path in enumerate(files):
        file_name = os.path.basename(file_path)
        file_stem = os.path.splitext(file_name)[0]

        # ===== 跳过已处理的文件 =====
        if file_name in done_source_files:
            print(f"\n  ⏭ [{idx+1}/{len(files)}] 跳过（已处理）: {file_name}")
            skip_count += 1
            continue

        print(f"\n{'━' * 70}")
        print(f"[{idx+1}/{len(files)}] 处理: {file_name}")
        print(f"{'━' * 70}")

        file_start = time.time()
        file_result = {
            "index": idx + 1,
            "source_file": file_name,
            "source_path": os.path.abspath(file_path),
            "status": "pending",
            "meeting_title": "",
            "output_files": {},
            "stats": {},
            "error": None,
        }

        try:
            # 1. 读取文本
            print(f"  读取文件...")
            raw_text = load_file_text(file_path)
            char_count = len(raw_text)
            print(f"  ✓ 读取完成: {char_count} 字")

            if char_count < 100:
                print(f"  ⚠ 文件内容过短（{char_count}字），跳过")
                file_result["status"] = "skipped"
                file_result["error"] = f"文件内容过短: {char_count}字"
                results_summary.append(file_result)
                fail_count += 1
                continue

            # 2. 提取元数据
            metadata = extract_metadata_from_text(raw_text)
            meeting_title = metadata.get("title", file_stem)
            if meeting_title in ("会议", ""):
                meeting_title = file_stem  # 用文件名作为标题
                metadata["title"] = meeting_title

            # 2.5 从文件名提取日期（优先于文本中提取的日期）
            fname_date = extract_date_from_filename(file_name)
            if fname_date:
                metadata["date"] = fname_date
                print(f"  日期（文件名）: {fname_date}")

            file_result["meeting_title"] = meeting_title
            print(f"  会议标题: {meeting_title}")

            # 拼接带日期的名称前缀（用于目录名和文件名）
            meeting_date = metadata.get("date", "")
            if meeting_date:
                name_prefix = f"{meeting_date}_{meeting_title}"
            else:
                name_prefix = meeting_title

            # 3. 创建该会议的输出子目录
            sub_dir_name = safe_filename(f"{idx+1:02d}_{name_prefix}")
            sub_dir = os.path.join(output_dir, sub_dir_name)
            os.makedirs(sub_dir, exist_ok=True)

            # 4. 生成纪要
            print(f"  生成会议纪要...")
            md_filename = safe_filename(name_prefix) + ".md"
            md_path = os.path.join(sub_dir, md_filename)

            result = generator.generate(
                raw_text=raw_text,
                metadata=metadata,
                output_path=md_path
            )

            file_result["output_files"]["md"] = md_path

            # 5. 生成 RAG 文档（复用上面的 result，不重复跑流水线）
            rag_docs = None
            if generate_rag:
                print(f"  生成RAG文档...")
                rag_docs = generator.generate_rag_documents(result=result)

                # 保存 RAG JSON
                rag_filename = safe_filename(name_prefix) + "_rag.json"
                rag_path = os.path.join(sub_dir, rag_filename)
                rag_output = {
                    "version": "1.0",
                    "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "source_file": file_name,
                    "meeting_title": meeting_title,
                    "meeting_date": metadata.get("date", ""),
                    "meeting_type": result.get("meeting_type", ""),
                    "total_documents": len(rag_docs),
                    "documents": rag_docs
                }
                with open(rag_path, 'w', encoding='utf-8') as f:
                    json.dump(rag_output, f, ensure_ascii=False, indent=2)
                file_result["output_files"]["rag_json"] = rag_path
                print(f"  ✓ RAG文档: {len(rag_docs)} 个块")

                # 收集到合并列表
                if merge_rag:
                    for doc in rag_docs:
                        doc["metadata"]["source_file"] = file_name
                    all_rag_docs.extend(rag_docs)

            # 6. 保存元信息
            meta_filename = safe_filename(name_prefix) + "_meta.json"
            meta_path = os.path.join(sub_dir, meta_filename)
            meta_data = {
                "source_file": file_name,
                "meeting_title": meeting_title,
                "metadata": result.get("metadata", {}),
                "speakers": result.get("speakers", {}),
                "topics": [{"id": t.get("id"), "title": t.get("title"),
                            "summary": t.get("summary", "")}
                           for t in result.get("topics", [])],
                "meeting_type": result.get("meeting_type", ""),
                "quality_report": result.get("quality_report", {}),
                "duration": result.get("duration", 0),
            }
            with open(meta_path, 'w', encoding='utf-8') as f:
                json.dump(meta_data, f, ensure_ascii=False, indent=2)
            file_result["output_files"]["meta_json"] = meta_path

            # 统计
            elapsed = time.time() - file_start
            quality_score = result.get("quality_report", {}).get("score", 0)
            topic_count = len(result.get("topics", []))
            md_chars = len(result.get("markdown", ""))

            # --- 判定是否降级 ---
            degraded_reasons = []
            if DEGRADED_REQUIRE_TOPICS and topic_count == 0:
                degraded_reasons.append("0 主题")
            if quality_score < DEGRADED_QUALITY_THRESHOLD:
                degraded_reasons.append(f"质量分 {quality_score} < {DEGRADED_QUALITY_THRESHOLD}")
            if md_chars < DEGRADED_MIN_OUTPUT_CHARS:
                degraded_reasons.append(f"纪要过短 {md_chars} 字")

            if degraded_reasons:
                file_result["status"] = "degraded"
                file_result["degraded_reasons"] = degraded_reasons
            else:
                file_result["status"] = "success"

            file_result["stats"] = {
                "source_chars": char_count,
                "output_chars": md_chars,
                "topics": topic_count,
                "quality_score": quality_score,
                "rag_docs": len(rag_docs) if rag_docs else 0,
                "duration_seconds": round(elapsed, 1),
            }

            if file_result["status"] == "success":
                success_count += 1
                print(f"  ✓ 完成! 耗时 {elapsed:.1f}s, "
                      f"质量 {quality_score}/100, "
                      f"{topic_count} 个主题, "
                      f"{md_chars} 字纪要")
            else:
                degraded_count += 1
                print(f"  ⚠ 降级完成! 耗时 {elapsed:.1f}s, "
                      f"质量 {quality_score}/100, "
                      f"{topic_count} 个主题 — 原因: {'; '.join(degraded_reasons)}")

        except Exception as e:
            elapsed = time.time() - file_start
            file_result["status"] = "failed"
            file_result["error"] = str(e)
            file_result["stats"]["duration_seconds"] = round(elapsed, 1)
            fail_count += 1
            print(f"  ✗ 失败: {e}")
            traceback.print_exc()

        results_summary.append(file_result)

    # ============================================================
    # 批量处理完成，输出汇总
    # ============================================================

    total_elapsed = time.time() - total_start

    # 合并所有 RAG 文档到一个文件
    merged_rag_path = None
    if merge_rag and all_rag_docs:
        merged_rag_path = os.path.join(output_dir, "all_rag_documents.json")
        merged = {
            "version": "1.0",
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "total_meetings": success_count,
            "total_documents": len(all_rag_docs),
            "documents": all_rag_docs,
        }
        with open(merged_rag_path, 'w', encoding='utf-8') as f:
            json.dump(merged, f, ensure_ascii=False, indent=2)

    # 保存批量处理汇总
    summary = {
        "batch_info": {
            "input_dir": os.path.abspath(input_dir),
            "output_dir": os.path.abspath(output_dir),
            "processed_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "total_files": len(files),
            "skipped": skip_count,
            "success": success_count,
            "degraded": degraded_count,
            "failed": fail_count,
            "total_duration_seconds": round(total_elapsed, 1),
        },
        "files": results_summary,
    }

    summary_path = os.path.join(output_dir, "batch_summary.json")
    with open(summary_path, 'w', encoding='utf-8') as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)

    # 打印汇总
    print(f"\n{'=' * 70}")
    print(f"批量处理完成")
    print(f"{'=' * 70}")
    print(f"  总文件:   {len(files)}")
    print(f"  跳过:     {skip_count}")
    print(f"  成功:     {success_count}")
    print(f"  降级:     {degraded_count}")
    print(f"  失败:     {fail_count}")
    print(f"  总耗时:   {total_elapsed:.1f} 秒")
    if success_count + degraded_count > 0:
        print(f"  平均耗时: {total_elapsed/(success_count + degraded_count):.1f} 秒/篇")
    print(f"\n  输出目录: {os.path.abspath(output_dir)}")
    print(f"  处理汇总: {summary_path}")
    if merged_rag_path:
        print(f"  合并RAG:  {merged_rag_path} ({len(all_rag_docs)} 个文档块)")

    print(f"\n  各文件结果:")
    for r in results_summary:
        icon_map = {"success": "✓", "degraded": "⚠", "failed": "✗", "skipped": "⊘"}
        status_icon = icon_map.get(r["status"], "?")
        stats = r.get("stats", {})
        line = f"    {status_icon} [{r['index']}] {r['source_file']}"
        if r["status"] in ("success", "degraded"):
            line += (f" → {r['meeting_title']} "
                     f"| {stats.get('quality_score', '?')}/100 "
                     f"| {stats.get('topics', '?')}主题 "
                     f"| {stats.get('duration_seconds', '?')}s")
            if r["status"] == "degraded":
                line += f" | ⚠ {'; '.join(r.get('degraded_reasons', []))}"
        elif r["status"] == "failed":
            line += f" → 错误: {r.get('error', '')}"
        else:
            line += f" → {r.get('error', '跳过')}"
        print(line)

    print(f"{'=' * 70}")

    return summary


# ============================================================
# 命令行入口
# ============================================================

def main():
    """命令行入口"""
    import argparse

    parser = argparse.ArgumentParser(
        description="批量会议纪要生成工具：从文件夹读取 .docx/.txt 文件，生成 .md 纪要和 .json RAG文档",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  python batch_meeting.py ./会议录音/               # 输出到 ./output
  python batch_meeting.py ./会议录音/ ./报告输出/   # 指定输出目录
  python batch_meeting.py ./input --no-rag          # 不生成RAG文档
  python batch_meeting.py ./input --no-merge         # 不合并RAG文档

输出结构：
  输出目录/
  ├── 01_XX会议/
  │   ├── XX会议.md              ← Markdown纪要
  │   ├── XX会议_rag.json        ← RAG文档块
  │   └── XX会议_meta.json       ← 元信息
  ├── 02_YY会议/
  │   └── ...
  ├── batch_summary.json          ← 批量处理汇总
  └── all_rag_documents.json      ← 所有RAG文档合并
        """
    )

    parser.add_argument("input_dir", nargs="?", default="./input",
                        help="输入文件夹路径（默认 ./input）")
    parser.add_argument("output_dir", nargs="?", default="./output",
                        help="输出目录（默认 ./output）")
    parser.add_argument("--no-rag", action="store_true",
                        help="不生成 RAG 文档")
    parser.add_argument("--no-merge", action="store_true",
                        help="不合并所有 RAG 文档到一个文件")

    args = parser.parse_args()

    if not os.path.isdir(args.input_dir):
        print(f"错误：输入目录不存在: {args.input_dir}")
        sys.exit(1)

    # 初始化 UTF-8 日志（必须在任何 print 之前完成，才能全部捕获到日志文件）
    setup_logging(args.output_dir)

    batch_process(
        input_dir=args.input_dir,
        output_dir=args.output_dir,
        generate_rag=not args.no_rag,
        merge_rag=not args.no_merge,
    )


if __name__ == "__main__":
    main()
