#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
会议纪要 Markdown → DOCX 排版转换器（python-docx 版）

与 pdf_renderer.py 的输出在视觉上保持一致：
  - 封面：蓝色标题居中，日期靠下，无横线、无副标题
  - 正文：四级标题分色、段首缩进、列表、引用块、加粗/斜体/代码
  - 页脚：页码居中（封面不计页码）

依赖：
  pip install python-docx

用法：
  python docx_renderer.py 会议纪要.md
  python docx_renderer.py 会议纪要.md -o ./out/会议纪要.docx
"""

import os
import re
from typing import Tuple

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================
# 字体与颜色（与 PDF 版保持一致）
# ============================================================

CN_FONT = "Microsoft YaHei"  # 中文字体；若系统无此字体，Word 自动 fallback

CLR_TITLE = RGBColor(0x1a, 0x52, 0x76)   # 封面标题：蓝色
CLR_H1    = RGBColor(0x1a, 0x1a, 0x2e)
CLR_H2    = RGBColor(0x1a, 0x52, 0x76)
CLR_H3    = RGBColor(0x2c, 0x3e, 0x50)
CLR_H4    = RGBColor(0x34, 0x49, 0x5e)
CLR_H5    = RGBColor(0x1f, 0x61, 0x8d)
CLR_BODY  = RGBColor(0x22, 0x22, 0x22)
CLR_QUOTE = RGBColor(0x66, 0x66, 0x66)
CLR_META  = RGBColor(0x44, 0x44, 0x44)


# ============================================================
# 小工具：设置字体（中英文都生效）+ 加 run
# ============================================================

def _set_font(run, size_pt: float, color: RGBColor = None, bold: bool = False,
              italic: bool = False, font_name: str = CN_FONT, mono: bool = False):
    name = "Courier New" if mono else font_name
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def _add_runs_from_markdown(paragraph, text: str, base_size: float = 11,
                            base_color: RGBColor = CLR_BODY, base_bold: bool = False):
    """解析行内 **bold** / *italic* / `code`，依次加 run 到段落。"""
    # 用占位符替换，避免正则相互吃字符
    parts = []
    pos = 0
    pattern = re.compile(
        r'(\*\*([^*\n]+?)\*\*)|(\*([^*\n]+?)\*)|(`([^`\n]+?)`)'
    )
    for m in pattern.finditer(text):
        if m.start() > pos:
            parts.append(("plain", text[pos:m.start()]))
        if m.group(1):
            parts.append(("bold", m.group(2)))
        elif m.group(3):
            parts.append(("italic", m.group(4)))
        elif m.group(5):
            parts.append(("code", m.group(6)))
        pos = m.end()
    if pos < len(text):
        parts.append(("plain", text[pos:]))
    if not parts:
        parts = [("plain", text)]

    for kind, content in parts:
        run = paragraph.add_run(content)
        if kind == "bold":
            _set_font(run, base_size, base_color, bold=True)
        elif kind == "italic":
            _set_font(run, base_size, base_color, italic=True, bold=base_bold)
        elif kind == "code":
            _set_font(run, base_size - 1, base_color, mono=True, bold=base_bold)
        else:
            _set_font(run, base_size, base_color, bold=base_bold)


def _add_hr(paragraph):
    """给段落加底边框，模拟分隔线。"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "dddddd")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_line_spacing(paragraph, leading_pt: float):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(leading_pt)


def _add_page_number_field(paragraph):
    """在段落里插入 Word 的 PAGE 字段。"""
    run = paragraph.add_run()
    _set_font(run, 9, RGBColor(0x99, 0x99, 0x99))

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE  \\* MERGEFORMAT"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _set_page_number_start(section, start: int = 1):
    """让某 section 的页码从 start 开始。"""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:start"), str(start))


# ============================================================
# 元信息提取（与 pdf_renderer 同源逻辑）
# ============================================================

def extract_meeting_info(md_text: str) -> dict:
    info = {"title": "会议纪要", "date": "", "meeting_type": ""}
    m = re.search(r'^#\s+(.+)$', md_text, re.MULTILINE)
    if m:
        t = re.sub(r'\*+', '', m.group(1)).strip()
        if t:
            info["title"] = t
    m = re.search(r'\*\*日期\*\*[：:]\s*(.+)', md_text)
    if m:
        info["date"] = m.group(1).strip()
    else:
        m = re.search(r'(\d{4}年\d{1,2}月\d{1,2}日)', md_text)
        if m:
            info["date"] = m.group(1)
    m = re.search(r'\*\*会议类型\*\*[：:]\s*(.+)', md_text)
    if m:
        info["meeting_type"] = m.group(1).strip()
    return info


_SUBHEAD_RE = re.compile(r'^\*\*([^*\n]{2,40})\*\*\s*[：:]?\s*$')


# ============================================================
# 主转换
# ============================================================

def md_to_docx(md_path: str, docx_path: str) -> str:
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    info = extract_meeting_info(md_text)
    doc = Document()

    # ---- 全局页边距（与 PDF 一致：上下 2cm，左右 2cm）----
    cover_section = doc.sections[0]
    cover_section.top_margin = Cm(2)
    cover_section.bottom_margin = Cm(2)
    cover_section.left_margin = Cm(2)
    cover_section.right_margin = Cm(2)
    # 封面不显示页码：用 different_first_page，封面页脚留空
    # 但封面是独立 section，更简单：直接不在封面 section 的 footer 写东西

    # ====== 封面 ======
    # 顶部空白：约 9cm（PDF 是 9cm Spacer）
    for _ in range(9):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _set_line_spacing(p, 18)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(10)
    _set_line_spacing(p_title, 40)
    run = p_title.add_run(info["title"])
    _set_font(run, 28, CLR_TITLE, bold=True)

    if info["date"]:
        # 标题与日期之间 4cm 留白
        for _ in range(5):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            _set_line_spacing(p, 18)
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_line_spacing(p_date, 36)
        run = p_date.add_run(info["date"])
        _set_font(run, 24, CLR_META, bold=True)

    # 封面 section 结束：插入分节符（下一页起新 section）
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2)
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2)
    # 让正文 section 的页码从 1 开始
    _set_page_number_start(new_section, 1)
    # 正文 section 的页脚：页码
    new_section.footer.is_linked_to_previous = False
    foot_p = new_section.footer.paragraphs[0]
    foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot_p.add_run("第 ")
    _set_font(foot_p.runs[0], 9, RGBColor(0x99, 0x99, 0x99))
    _add_page_number_field(foot_p)
    run_tail = foot_p.add_run(" 页")
    _set_font(run_tail, 9, RGBColor(0x99, 0x99, 0x99))
    # 封面 section 的页脚保持空
    cover_section.footer.is_linked_to_previous = False

    # ====== 正文解析 ======
    lines = md_text.split("\n")
    i = 0
    skip_header_block = True
    subhead_idx = 0  # 当前 H2 章节下的小标题序号，遇到新 H2 重置

    def add_heading(text: str, size: float, color: RGBColor, *,
                    align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=4,
                    leading=20, indent_left: float = 0):
        p = doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        if indent_left:
            pf.left_indent = Pt(indent_left)
        _set_line_spacing(p, leading)
        _add_runs_from_markdown(p, text, base_size=size, base_color=color, base_bold=True)
        return p

    def add_body(text: str, indent_first: bool = True):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        if indent_first:
            pf.first_line_indent = Cm(0.74)  # ≈ 22pt
        _set_line_spacing(p, 20)
        _add_runs_from_markdown(p, text, base_size=11, base_color=CLR_BODY)
        return p

    def add_list_item(text: str, bullet: str = "•  "):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Pt(24)
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        _set_line_spacing(p, 20)
        run_b = p.add_run(bullet)
        _set_font(run_b, 11, CLR_BODY)
        _add_runs_from_markdown(p, text, base_size=11, base_color=CLR_BODY)
        return p

    def add_quote(text: str):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Pt(16)
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        _set_line_spacing(p, 16)
        _add_runs_from_markdown(p, text, base_size=10, base_color=CLR_QUOTE)
        return p

    def add_subhead(label: str):
        # label 由调用方传入完整文案（已含 "1、" 之类的前缀），此处不再追加装饰
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(2)
        _set_line_spacing(p, 18)
        run = p.add_run(label)
        _set_font(run, 11, CLR_H5, bold=True)
        return p

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 跳过幻觉提示块
        if stripped.startswith("> ⚠"):
            i += 1
            while i < len(lines) and lines[i].strip().startswith(">"):
                i += 1
            continue

        # 跳过开头元信息区
        if skip_header_block:
            if stripped.startswith("## "):
                h2_text = stripped[3:].strip()
                if "目录" in h2_text:
                    i += 1
                    while i < len(lines):
                        ns = lines[i].strip()
                        if ns.startswith("## ") or ns.startswith("# "):
                            break
                        i += 1
                    continue
                skip_header_block = False
            else:
                i += 1
                continue

        if not stripped:
            i += 1
            continue

        # h1
        if stripped.startswith("# ") and not stripped.startswith("## "):
            text = stripped[2:].strip()
            p = add_heading(text, 20, CLR_H1,
                            align=WD_ALIGN_PARAGRAPH.CENTER,
                            space_before=10, space_after=8, leading=28)
            # h1 下方一条粗一点的分割线
            hr_p = doc.add_paragraph()
            _set_line_spacing(hr_p, 6)
            _add_hr(hr_p)
            i += 1
            continue

        # h2
        if stripped.startswith("## "):
            text = stripped[3:].strip()
            if "目录" in text:
                i += 1
                while i < len(lines):
                    ns = lines[i].strip()
                    if ns.startswith("## ") or ns.startswith("# "):
                        break
                    i += 1
                continue
            subhead_idx = 0  # 进入新 H2 章节，小标题编号重置
            add_heading(text, 15, CLR_H2,
                        space_before=16, space_after=6, leading=24, indent_left=0)
            i += 1
            continue

        # h3
        if stripped.startswith("### "):
            text = stripped[4:].strip()
            add_heading(text, 12.5, CLR_H3, space_before=10, space_after=4, leading=20)
            i += 1
            continue

        # h4
        if stripped.startswith("#### "):
            text = stripped[5:].strip()
            add_heading(text, 11.5, CLR_H4, space_before=8, space_after=3, leading=18)
            i += 1
            continue

        # 引用
        if stripped.startswith("> "):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith("> "):
                quote_lines.append(lines[i].strip()[2:].strip())
                i += 1
            add_quote(" ".join(quote_lines))
            continue

        # 无序列表
        if re.match(r'^[-*]\s+', stripped):
            text = re.sub(r'^[-*]\s+', '', stripped)
            if not text.strip() and i + 1 < len(lines) and lines[i + 1].strip():
                next_s = lines[i + 1].strip()
                if not next_s.startswith('#') and not re.match(r'^[-*]\s+', next_s) \
                        and not re.match(r'^\d+[.、]', next_s):
                    text = next_s
                    i += 1
            if text.strip():
                add_list_item(text, bullet="•  ")
            i += 1
            continue

        # 有序列表
        m = re.match(r'^(\d+)[.、]\s*(.*)', stripped)
        if m:
            num = m.group(1)
            text = m.group(2)
            if not text.strip() and i + 1 < len(lines) and lines[i + 1].strip():
                next_s = lines[i + 1].strip()
                if not next_s.startswith('#') and not re.match(r'^[-*]\s+', next_s) \
                        and not re.match(r'^\d+[.、]', next_s):
                    text = next_s
                    i += 1
            add_list_item(text, bullet=f"{num}. ")
            i += 1
            continue

        # 分隔线
        if re.match(r'^[-*_]{3,}\s*$', stripped):
            hr_p = doc.add_paragraph()
            _set_line_spacing(hr_p, 6)
            _add_hr(hr_p)
            i += 1
            continue

        # 目录式行（[文本](#锚)）
        if re.match(r'^\d+\.\s*\[.+\]\(#.+\)$', stripped):
            i += 1
            continue

        # 段首加粗标签 → 小标题
        sub_m = _SUBHEAD_RE.match(stripped)
        if sub_m:
            subhead_idx += 1
            add_subhead(f"{subhead_idx}、{sub_m.group(1).strip()}")
            i += 1
            continue

        # 普通段落（合并连续非空行）
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_stripped = lines[i].strip()
            if not next_stripped or next_stripped.startswith("#") or \
               next_stripped.startswith("> ") or next_stripped.startswith("- ") or \
               next_stripped.startswith("* ") or re.match(r'^\d+[.、]', next_stripped) or \
               re.match(r'^[-*_]{3,}', next_stripped) or _SUBHEAD_RE.match(next_stripped):
                break
            para_lines.append(next_stripped)
            i += 1

        raw = " ".join(para_lines)
        inline_sub = re.match(r'^\*\*([^*\n]{2,40})\*\*\s*[：:]\s*(.+)$', raw)
        if inline_sub:
            subhead_idx += 1
            add_subhead(f"{subhead_idx}、{inline_sub.group(1).strip()}")
            add_body(inline_sub.group(2).strip())
        else:
            add_body(raw)

    os.makedirs(os.path.dirname(docx_path) or ".", exist_ok=True)
    doc.save(docx_path)
    return docx_path


